"""
Word (.docx) session report generator for VNS-TA.

Generates a clinical-grade session summary document containing
pre/post vitals, treatment parameters, intra-session statistics,
annotations, and clinician notes.
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from vns_ta.wizard_pages import EX4_CHANNEL_PARAMS


def _add_heading(doc: Document, text: str, level: int = 2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]):
    """Two-column table: label | value."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        cell = table.rows[i].cells[1]
        cell.text = value
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    doc.add_paragraph("")


def _fmt(val, unit: str = "", precision: int = 1) -> str:
    if val is None:
        return "--"
    if isinstance(val, float):
        return f"{val:.{precision}f} {unit}".strip()
    return f"{val} {unit}".strip()


def generate_session_report(path: str, data: dict) -> None:
    """Generate a .docx session report.

    Expected keys in *data*:
        baseline_hr, baseline_rmssd, spo2,
        last_hr, last_rmssd,
        modality_name, active_channels,
        session_start (datetime), session_end (datetime),
        csv_path,
        annotations  -- list of (timestamp_str, text)
        hr_values, rmssd_values  -- lists of floats for stats
        notes  -- str from clinician
        checklist  -- list of (str, bool)
        outcome  -- "normal" | "aborted"
    """
    doc = Document()

    # Title
    title = doc.add_heading("VNS-TA Session Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    now = data.get("session_end") or datetime.now()
    start = data.get("session_start")
    duration_min = "--"
    if start and now:
        duration_min = f"{(now - start).total_seconds() / 60:.1f}"

    date_str = now.strftime("%A, %B %d, %Y  %H:%M")
    doc.add_paragraph(f"Date: {date_str}")

    # Section 1: Session Overview
    _add_heading(doc, "Session Overview")
    outcome = data.get("outcome", "normal")
    outcome_text = ("Session completed normally."
                    if outcome == "normal"
                    else "Session was ABORTED.")
    _add_key_value_table(doc, [
        ("Outcome", outcome_text),
        ("Modality", data.get("modality_name", "--")),
        ("Active Channels", ", ".join(
            f"CH {c}" for c in data.get("active_channels", [])
        ) or "--"),
        ("Total Duration", f"{duration_min} minutes"),
    ])

    # Section 2: Pre-Session Baselines
    _add_heading(doc, "Pre-Session Baselines")
    _add_key_value_table(doc, [
        ("Heart Rate", _fmt(data.get("baseline_hr"), "bpm", 0)),
        ("RMSSD", _fmt(data.get("baseline_rmssd"), "ms")),
        ("SpO\u2082", _fmt(data.get("spo2"), "%", 0)),
    ])

    # Section 3: Post-Session Readings
    _add_heading(doc, "Post-Session Readings")
    pre_rmssd = data.get("baseline_rmssd")
    post_rmssd = data.get("last_rmssd")
    delta_rmssd = "--"
    if pre_rmssd and post_rmssd and pre_rmssd > 0:
        pct = ((post_rmssd - pre_rmssd) / pre_rmssd) * 100
        delta_rmssd = f"{pct:+.1f}%"
    _add_key_value_table(doc, [
        ("Heart Rate", _fmt(data.get("last_hr"), "bpm", 0)),
        ("RMSSD", _fmt(post_rmssd, "ms")),
        ("\u0394 RMSSD from Baseline", delta_rmssd),
    ])

    # Section 4: Treatment Parameters
    _add_heading(doc, "Treatment Parameters")
    active = data.get("active_channels", [])
    params = [p for p in EX4_CHANNEL_PARAMS if p["ch"] in active]
    if params:
        headers = ["Channel", "Target", "Frequency",
                   "Pulse Width", "Duty Cycle", "Intensity"]
        table = doc.add_table(rows=1 + len(params), cols=len(headers))
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for r, p in enumerate(params, start=1):
            table.rows[r].cells[0].text = f"CH {p['ch']}"
            table.rows[r].cells[1].text = p["target"]
            table.rows[r].cells[2].text = p["frequency"]
            table.rows[r].cells[3].text = p["pulse_width"]
            table.rows[r].cells[4].text = p["duty_cycle"]
            table.rows[r].cells[5].text = p["intensity"]
        doc.add_paragraph("")
    else:
        doc.add_paragraph("No channel parameters available.")

    # Section 5: Intra-Session Statistics
    _add_heading(doc, "Intra-Session Statistics")
    hr_vals = data.get("hr_values", [])
    rmssd_vals = data.get("rmssd_values", [])
    stats_rows = []
    if hr_vals:
        stats_rows.extend([
            ("HR Min", f"{min(hr_vals):.0f} bpm"),
            ("HR Max", f"{max(hr_vals):.0f} bpm"),
            ("HR Avg", f"{sum(hr_vals)/len(hr_vals):.0f} bpm"),
        ])
    else:
        stats_rows.append(("Heart Rate", "No data recorded"))
    if rmssd_vals:
        stats_rows.extend([
            ("RMSSD Min", f"{min(rmssd_vals):.1f} ms"),
            ("RMSSD Max", f"{max(rmssd_vals):.1f} ms"),
            ("RMSSD Avg", f"{sum(rmssd_vals)/len(rmssd_vals):.1f} ms"),
        ])
    else:
        stats_rows.append(("RMSSD", "No data recorded"))
    _add_key_value_table(doc, stats_rows)

    # Section 6: Annotations
    annotations = data.get("annotations", [])
    if annotations:
        _add_heading(doc, "Session Annotations")
        table = doc.add_table(rows=1 + len(annotations), cols=2)
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "Time"
        table.rows[0].cells[1].text = "Annotation"
        for r_p in table.rows[0].cells:
            for p in r_p.paragraphs:
                for run in p.runs:
                    run.bold = True
        for r, (ts, txt) in enumerate(annotations, start=1):
            table.rows[r].cells[0].text = ts
            table.rows[r].cells[1].text = txt
        doc.add_paragraph("")

    # Section 7: Clinician Notes
    notes = data.get("notes", "").strip()
    _add_heading(doc, "Clinician Notes")
    doc.add_paragraph(notes if notes else "(No notes entered.)")

    # Section 8: Post-Session Checklist
    _add_heading(doc, "Post-Session Checklist")
    checklist = data.get("checklist", [])
    if checklist:
        for item_text, checked in checklist:
            mark = "\u2611" if checked else "\u2610"
            doc.add_paragraph(f"{mark}  {item_text}")
    else:
        doc.add_paragraph("(Checklist not available.)")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated by VNS-TA on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.save(path)
