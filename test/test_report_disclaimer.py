from __future__ import annotations

from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from hnh.report import generate_session_report


class ReportDisclaimerTests(unittest.TestCase):
    @staticmethod
    def _doc_table_text(doc: Document) -> str:
        chunks: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        return "\n".join(chunks)

    def test_report_includes_legal_disclaimer_section_when_disclaimer_provided(self):
        with TemporaryDirectory() as tmp:
            report_path = Path(tmp) / "session_report.docx"
            data = {
                "session_id": "s-001",
                "profile_id": "Admin",
                "session_type": "General Monitoring",
                "session_start": datetime(2026, 2, 25, 9, 0, 0),
                "session_end": datetime(2026, 2, 25, 9, 15, 0),
                "baseline_hr": 72.0,
                "baseline_rmssd": 28.5,
                "last_hr": 68.0,
                "last_rmssd": 35.1,
                "annotations": [("09:05:00", "Steady breathing")],
                "hr_values": [72.0, 70.0, 68.0],
                "rmssd_values": [28.5, 31.0, 35.1],
                "notes": "Draft test note.",
                "csv_path": "session.csv",
                "report_stage": "final",
                "qtc": {},
                "disclaimer": {
                    "warning": "RESEARCH USE ONLY - NOT FOR CLINICAL DIAGNOSIS OR TREATMENT.",
                    "source_path": "hnh/disclaimer.md",
                    "text": "# Research Use Disclaimer (Full Text)\n\nSample disclaimer body.",
                    "sha256": "abc123",
                    "acknowledgment_mode": "interactive_dialog",
                    "acknowledged_at": "2026-02-25T09:00:00",
                },
            }

            generate_session_report(str(report_path), data)

            self.assertTrue(report_path.exists())
            doc = Document(str(report_path))
            all_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Legal Disclaimer", all_text)
            self.assertIn(
                "RESEARCH USE ONLY - NOT FOR CLINICAL DIAGNOSIS OR TREATMENT.",
                all_text,
            )
            self.assertIn("Saved Disclaimer Text", all_text)
            self.assertIn("Sample disclaimer body.", all_text)

    def test_report_includes_annotation_associations_section_when_present(self):
        with TemporaryDirectory() as tmp:
            report_path = Path(tmp) / "session_report.docx"
            data = {
                "session_id": "s-001",
                "profile_id": "Admin",
                "session_type": "General Monitoring",
                "session_start": datetime(2026, 2, 25, 9, 0, 0),
                "session_end": datetime(2026, 2, 25, 9, 15, 0),
                "baseline_hr": 72.0,
                "baseline_rmssd": 28.5,
                "last_hr": 68.0,
                "last_rmssd": 35.1,
                "annotations": [("09:05:00", "Steady breathing")],
                "annotation_associations": [
                    {
                        "annotation": "Caffeine intake",
                        "events": 8,
                        "sessions": 4,
                        "delta_hr_bpm": 5.4,
                        "delta_rmssd_ms": -3.2,
                        "confidence": "Moderate",
                        "caveat": "small sample",
                    }
                ],
                "annotation_associations_method": "Method: pre window t-45s..t-5s; post window t+5s..t+45s.",
                "hr_values": [72.0, 70.0, 68.0],
                "rmssd_values": [28.5, 31.0, 35.1],
                "notes": "Draft test note.",
                "csv_path": "session.csv",
                "report_stage": "final",
                "qtc": {},
                "disclaimer": {},
            }

            generate_session_report(str(report_path), data)

            self.assertTrue(report_path.exists())
            doc = Document(str(report_path))
            all_text = "\n".join(p.text for p in doc.paragraphs)
            table_text = self._doc_table_text(doc)
            self.assertIn("Annotation Associations (Exploratory)", all_text)
            self.assertIn("pre window t-45s..t-5s", all_text)
            self.assertIn("Caffeine intake", table_text)
            self.assertIn("Moderate", table_text)

if __name__ == "__main__":
    unittest.main()
