"""
Word (.docx) session report generator for Hertz & Hearts.

Generates a session summary document containing pre/post vitals,
intra-session statistics, annotations, and user notes.
"""

from datetime import datetime
import locale
from pathlib import Path
from typing import Any

import numpy as np
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_heading(doc: Document, text: str, level: int = 2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    return h


def _date_for_region(dt: datetime) -> str:
    try:
        locale.setlocale(locale.LC_TIME, "")
    except locale.Error:
        pass
    try:
        # Prefer OS locale date rendering when available.
        locale_date = dt.strftime("%x")
        if locale_date and "%" not in locale_date:
            return locale_date
    except Exception:
        pass
    # Safe default/fallback: US date format.
    return dt.strftime("%m/%d/%Y")


def _uses_24_hour_time() -> bool:
    try:
        locale.setlocale(locale.LC_TIME, "")
    except locale.Error:
        pass
    try:
        tfmt = locale.nl_langinfo(locale.T_FMT)
        if "%H" in tfmt:
            return True
        if "%I" in tfmt:
            return False
    except Exception:
        pass
    loc = ""
    try:
        loc = (locale.getlocale(locale.LC_TIME)[0] or "").lower()
    except Exception:
        loc = ""
    verified_24h = {
        "en_gb", "en_ie", "de_de", "fr_fr", "es_es", "it_it", "nl_nl",
        "sv_se", "fi_fi", "no_no", "da_dk", "pl_pl", "cs_cz", "sk_sk",
        "hu_hu", "ro_ro", "tr_tr", "pt_pt", "ru_ru", "uk_ua", "zh_cn",
        "zh_tw", "ja_jp", "ko_kr",
    }
    return loc in verified_24h


def _report_datetime(value: datetime | str | None) -> str:
    if isinstance(value, datetime):
        dt = value
    elif isinstance(value, str):
        text = value.strip()
        if not text:
            return "--"
        try:
            dt = datetime.fromisoformat(text.replace("Z", "+00:00"))
        except ValueError:
            return text
    else:
        return "--"
    if _uses_24_hour_time():
        time_str = dt.strftime("%H:%M")
    else:
        time_str = dt.strftime("%I:%M %p").lstrip("0")
    return f"{_date_for_region(dt)} @ {time_str}"


def _append_page_field(paragraph):
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    field_run = paragraph.add_run()
    field_run.font.size = Pt(8)
    field_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    r = field_run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)


def _add_key_value_table(
    doc: Document,
    rows: list[tuple[str, str]],
    *,
    label_width_in: float = 2.6,
    value_width_in: float = 4.1,
):
    """Two-column table: label | value."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    label_w = Inches(label_width_in)
    value_w = Inches(value_width_in)
    for i, (label, value) in enumerate(rows):
        left = table.rows[i].cells[0]
        right = table.rows[i].cells[1]
        left.width = label_w
        right.width = value_w
        left.text = label
        cell = right
        cell.text = value
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    doc.add_paragraph("")


def _fmt(val, unit: str = "", precision: int = 1) -> str:
    if val is None:
        return "--"
    if isinstance(val, float):
        return f"{val:.{precision}f} {unit}".strip()
    return f"{val} {unit}".strip()


def _fmt_qtc_session_value(qtc_data: dict) -> str:
    value = qtc_data.get("session_value_ms")
    if value is None:
        quality = qtc_data.get("quality", {})
        reason = quality.get("reason")
        if reason:
            return f"QTc unavailable ({reason})"
        return "QTc unavailable (signal quality too low)"
    return _fmt(value, "ms", 0)


def _to_float(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _derive_recovery_state(data: dict) -> tuple[str, list[str]]:
    baseline_rmssd = _to_float(data.get("baseline_rmssd"))
    last_rmssd = _to_float(data.get("last_rmssd"))
    baseline_hr = _to_float(data.get("baseline_hr"))
    last_hr = _to_float(data.get("last_hr"))

    score = 0
    drivers: list[str] = []

    if baseline_rmssd is not None and last_rmssd is not None:
        if last_rmssd >= baseline_rmssd:
            score += 1
            drivers.append("RMSSD is at/above baseline.")
        else:
            score -= 1
            drivers.append("RMSSD is below baseline.")
    else:
        drivers.append("RMSSD baseline comparison unavailable.")

    if baseline_hr is not None and last_hr is not None:
        if last_hr <= baseline_hr:
            score += 1
            drivers.append("Heart rate is at/below baseline.")
        else:
            score -= 1
            drivers.append("Heart rate is above baseline.")
    else:
        drivers.append("Heart-rate baseline comparison unavailable.")

    if score >= 2:
        return "High", drivers[:3]
    if score <= -1:
        return "Low", drivers[:3]
    return "Moderate", drivers[:3]


def _add_image_with_caption(
    doc: Document,
    caption: str,
    image_path: Path,
    width_inches: float = 6.8,
    *,
    keep_block: bool = True,
):
    """Insert caption + image with optional keep-together behavior."""
    caption_para = doc.add_paragraph(caption.strip())
    caption_para.paragraph_format.keep_with_next = keep_block
    image_para = doc.add_paragraph()
    image_para.paragraph_format.keep_together = keep_block
    image_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    doc.add_paragraph("")


def _build_visual_images(data: dict[str, Any], output_dir: Path) -> dict[str, Path]:
    """Generate trend/ecg PNGs and return available image paths."""
    visuals: dict[str, Path] = {}
    try:
        from matplotlib.figure import Figure
        from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
    except Exception:
        return visuals

    output_dir.mkdir(parents=True, exist_ok=True)
    hr_values = [float(v) for v in (data.get("hr_values") or []) if v is not None]
    rmssd_values = [float(v) for v in (data.get("rmssd_values") or []) if v is not None]
    hr_time_seconds = [float(v) for v in (data.get("hr_time_seconds") or []) if v is not None]
    rmssd_time_seconds = [float(v) for v in (data.get("rmssd_time_seconds") or []) if v is not None]
    ecg_values = [float(v) for v in (data.get("ecg_samples") or []) if v is not None]
    ecg_rate = int(data.get("ecg_sample_rate_hz") or 130)

    session_start = data.get("session_start")
    session_end = data.get("session_end")
    session_duration_sec = None
    if isinstance(session_start, datetime) and isinstance(session_end, datetime):
        session_duration_sec = max(1.0, (session_end - session_start).total_seconds())

    def _time_axis_minutes(values_len: int, explicit_seconds: list[float]) -> list[float]:
        if values_len <= 0:
            return []
        if len(explicit_seconds) == values_len:
            return [max(0.0, t) / 60.0 for t in explicit_seconds]
        if values_len == 1:
            return [0.0]
        if session_duration_sec is not None:
            return [float(t) / 60.0 for t in np.linspace(0.0, session_duration_sec, num=values_len)]
        return [float(i) for i in range(1, values_len + 1)]

    if hr_values or rmssd_values:
        trend_path = output_dir / "session_visual_trend.png"
        fig = Figure(figsize=(7.2, 2.4), dpi=140)
        canvas = FigureCanvas(fig)
        ax_hr = fig.add_subplot(211)
        ax_rmssd = fig.add_subplot(212)
        if hr_values:
            x_hr = _time_axis_minutes(len(hr_values), hr_time_seconds)
            ax_hr.plot(x_hr, hr_values, color="#C2185B", linewidth=1.6)
            ax_hr.set_ylabel("HR (bpm)", fontsize=8)
            ax_hr.grid(alpha=0.25)
        else:
            ax_hr.text(0.02, 0.5, "No HR trend data", transform=ax_hr.transAxes, fontsize=8)
        if rmssd_values:
            x_rmssd = _time_axis_minutes(len(rmssd_values), rmssd_time_seconds)
            ax_rmssd.plot(x_rmssd, rmssd_values, color="#1565C0", linewidth=1.6)
            ax_rmssd.set_ylabel("RMSSD (ms)", fontsize=8)
            ax_rmssd.set_xlabel("Elapsed session time (min)", fontsize=8)
            ax_rmssd.grid(alpha=0.25)
        else:
            ax_rmssd.text(0.02, 0.5, "No RMSSD trend data", transform=ax_rmssd.transAxes, fontsize=8)
        for axis in (ax_hr, ax_rmssd):
            axis.tick_params(axis="both", labelsize=7)
        fig.tight_layout(h_pad=0.6)
        canvas.print_png(str(trend_path))
        visuals["trend"] = trend_path

    if len(ecg_values) >= 20:
        strip_path = output_dir / "session_visual_ecg_strip.png"
        window = min(len(ecg_values), max(ecg_rate * 6, 260))
        selected = ecg_values[-window:]
        fig = Figure(figsize=(7.2, 1.7), dpi=140)
        canvas = FigureCanvas(fig)
        ax = fig.add_subplot(111)
        session_start = data.get("session_start")
        session_end = data.get("session_end")
        if isinstance(session_start, datetime) and isinstance(session_end, datetime):
            total_sec = max(0.0, (session_end - session_start).total_seconds())
        else:
            total_sec = len(ecg_values) / float(ecg_rate)
        strip_span_sec = len(selected) / float(ecg_rate)
        strip_start_sec = max(0.0, total_sec - strip_span_sec)
        t = [strip_start_sec + (idx / float(ecg_rate)) for idx in range(len(selected))]
        ax.plot(t, selected, color="#111111", linewidth=1.0)
        ax.set_title("Selected ECG strip (latest)", fontsize=9)
        ax.set_xlabel("Elapsed session time (s)", fontsize=8)
        ax.set_ylabel("Amplitude (mV)", fontsize=8)
        ax.grid(alpha=0.28)
        ax.tick_params(axis="both", labelsize=7)
        fig.tight_layout()
        canvas.print_png(str(strip_path))
        visuals["ecg"] = strip_path

    return visuals


def generate_session_share_pdf(path: str, data: dict) -> None:
    """Generate a one-page PDF summary from the same report data model."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    visuals = _build_visual_images(data, output.parent)

    title_style = ParagraphStyle(
        "ShareTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        textColor=colors.HexColor("#1A5276"),
        spaceAfter=6,
        alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        "ShareHeading",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#1A5276"),
        spaceBefore=5,
        spaceAfter=3,
    )
    body_style = ParagraphStyle(
        "ShareBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#222222"),
    )
    muted_style = ParagraphStyle(
        "ShareMuted",
        parent=body_style,
        fontSize=8,
        textColor=colors.HexColor("#6B7280"),
    )
    generated_style = ParagraphStyle(
        "ShareGenerated",
        parent=muted_style,
        alignment=TA_CENTER,
    )
    notice_style = ParagraphStyle(
        "ShareNotice",
        parent=body_style,
        textColor=colors.HexColor("#9C2727"),
        fontName="Helvetica-Bold",
    )

    session_end = data.get("session_end") or datetime.now()
    if not isinstance(session_end, datetime):
        session_end = datetime.now()

    start = data.get("session_start")
    if isinstance(start, datetime):
        duration_min = max((session_end - start).total_seconds() / 60.0, 0.0)
        duration_text = f"{duration_min:.1f} min"
    else:
        duration_text = "--"

    state, drivers = _derive_recovery_state(data)
    qtc_data = data.get("qtc", {}) or {}
    disclaimer = data.get("disclaimer", {}) or {}
    warning = str(disclaimer.get("warning", "")).strip()

    metrics_table = Table(
        [
            ["Metric", "Value"],
            ["HR (baseline/latest)", f"{_fmt(data.get('baseline_hr'), 'bpm', 0)} / {_fmt(data.get('last_hr'), 'bpm', 0)}"],
            ["RMSSD (baseline/latest)", f"{_fmt(data.get('baseline_rmssd'), 'ms')} / {_fmt(data.get('last_rmssd'), 'ms')}"],
            ["QTc (session average)", _fmt_qtc_session_value(qtc_data)],
        ],
        colWidths=[58 * mm, 122 * mm],
        repeatRows=1,
        hAlign="LEFT",
    )
    metrics_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1A5276")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D5D8DC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )

    profile_name = str(data.get("profile_id", "--"))
    prefix_story = [
        Paragraph(f"Hertz & Hearts - One-Page Session Report: {profile_name}", title_style),
        Paragraph(f"Report generated: {_report_datetime(session_end)}", generated_style),
        Spacer(1, 3 * mm),
        Paragraph("Session Overview", heading_style),
        Paragraph(
            f"Session Start Date & Time: {_report_datetime(start)}",
            body_style,
        ),
        Paragraph(f"Session Type: {data.get('session_type', 'General Monitoring')}", body_style),
        Paragraph(f"Duration: {duration_text}", body_style),
        Paragraph(f"Report Stage: {str(data.get('report_stage', 'final')).capitalize()}", body_style),
        Spacer(1, 3 * mm),
        Paragraph("Recovery Snapshot", heading_style),
        Paragraph(f"Recovery State: <b>{state}</b>", body_style),
    ]
    for driver in drivers:
        prefix_story.append(Paragraph(f"- {driver}", body_style))

    prefix_story.extend(
        [
            Spacer(1, 3 * mm),
            Paragraph("Core Metrics", heading_style),
            metrics_table,
        ]
    )

    suffix_story = []
    if warning:
        suffix_story.extend(
            [
                Spacer(1, 3 * mm),
                Paragraph("Research Notice", heading_style),
                Paragraph(warning, notice_style),
            ]
        )

    suffix_story.extend([Spacer(1, 4 * mm)])

    doc = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title="Hertz & Hearts One-Page Session Share",
    )

    def _flowable_height(flowable, avail_w: float, avail_h: float) -> float:
        _, h = flowable.wrap(avail_w, avail_h)
        before = flowable.getSpaceBefore() if hasattr(flowable, "getSpaceBefore") else 0.0
        after = flowable.getSpaceAfter() if hasattr(flowable, "getSpaceAfter") else 0.0
        return float(h + before + after)

    def _story_height(flowables: list, avail_w: float, avail_h: float) -> float:
        total = 0.0
        for fl in flowables:
            total += _flowable_height(fl, avail_w, avail_h)
        return total

    def _visual_block(mode: str) -> list:
        if not visuals:
            return []
        trend_path = visuals.get("trend")
        ecg_path = visuals.get("ecg")
        block: list = []
        if mode in {"both", "trend"} and trend_path and trend_path.exists():
            block.extend(
                [
                    Spacer(1, 3 * mm),
                    Paragraph("Session Visuals", heading_style),
                ]
            )
            if ecg_path and ecg_path.exists():
                block.extend(
                    [
                        Image(str(ecg_path), width=178 * mm, height=28 * mm),
                        Spacer(1, 2 * mm),
                    ]
                )
            block.extend(
                [
                    Image(str(trend_path), width=178 * mm, height=42 * mm),
                    Spacer(1, 2 * mm),
                ]
            )
            return block
        return []

    avail_w = A4[0] - doc.leftMargin - doc.rightMargin
    avail_h = A4[1] - doc.topMargin - doc.bottomMargin
    safety = 4 * mm

    core_height = _story_height(prefix_story + suffix_story, avail_w, avail_h)
    final_story = prefix_story + suffix_story

    # Keep one page with readable text by progressively reducing visual density.
    if visuals and core_height < (avail_h - safety):
        for mode in ("both", "trend"):
            candidate = prefix_story + _visual_block(mode) + suffix_story
            if _story_height(candidate, avail_w, avail_h) <= (avail_h - safety):
                final_story = candidate
                break
        else:
            final_story = prefix_story + [
                Spacer(1, 2 * mm),
                Paragraph("Visuals omitted in PDF to preserve one-page readability.", muted_style),
            ] + suffix_story

    doc.build(final_story)


def generate_session_report(path: str, data: dict) -> None:
    """Generate a .docx session report.

    Expected keys in *data*:
        baseline_hr, baseline_rmssd,
        last_hr, last_rmssd,
        qtc,
        session_type,
        session_start (datetime), session_end (datetime),
        csv_path,
        annotations  -- list of (timestamp_str, text)
        hr_values, rmssd_values  -- lists of floats for stats
        notes  -- str from user
        disclaimer -- dict with warning/text/source/hash/ack metadata
    """
    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    visuals = _build_visual_images(data, output.parent)
    doc = Document()

    # Title
    profile_name = str(data.get("profile_id", "--")).strip() or "Unknown"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = f"Hertz & Hearts Session Report: {profile_name}"
    title_run = title.add_run(title_text)
    title_size = 18
    if len(title_text) > 70:
        title_size = 14
    elif len(title_text) > 52:
        title_size = 16
    title_run.font.size = Pt(title_size)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    now = data.get("session_end") or datetime.now()
    start = data.get("session_start")
    duration_min = "--"
    if start and now:
        duration_min = f"{(now - start).total_seconds() / 60:.1f}"

    report_date = doc.add_paragraph(f"Report generated: {_report_datetime(now)}")
    report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 1: Session Overview
    _add_heading(doc, "Session Overview")
    session_type = data.get("session_type", "General Monitoring")
    report_stage = data.get("report_stage", "final").strip().lower()
    report_label = "Draft (in-progress)" if report_stage == "draft" else "Final"
    _add_key_value_table(doc, [
        (
            "Session Start Date & Time",
            _report_datetime(start),
        ),
        ("Session Type", session_type),
        ("Report Stage", report_label),
        ("Total Duration", f"{duration_min} minutes"),
    ], label_width_in=1.92, value_width_in=3.04)

    # Section 2: Pre-Session Baselines
    _add_heading(doc, "Pre-Session Baselines")
    _add_key_value_table(doc, [
        ("Heart Rate", _fmt(data.get("baseline_hr"), "bpm", 0)),
        ("RMSSD", _fmt(data.get("baseline_rmssd"), "ms")),
    ], label_width_in=1.92, value_width_in=3.04)

    # Section 3: Post-Session Readings
    _add_heading(doc, "Post-Session Readings")
    pre_rmssd = data.get("baseline_rmssd")
    post_rmssd = data.get("last_rmssd")
    qtc_data = data.get("qtc", {}) or {}
    delta_rmssd = "--"
    if pre_rmssd is not None and post_rmssd is not None and pre_rmssd > 0:
        pct = ((post_rmssd - pre_rmssd) / pre_rmssd) * 100
        delta_rmssd = f"{pct:+.1f}%"
    _add_key_value_table(doc, [
        ("Heart Rate", _fmt(data.get("last_hr"), "bpm", 0)),
        ("RMSSD", _fmt(post_rmssd, "ms")),
        ("\u0394 RMSSD from Baseline", delta_rmssd),
        ("QTc (session average)", _fmt_qtc_session_value(qtc_data)),
    ], label_width_in=1.92, value_width_in=3.04)
    qtc_trend = qtc_data.get("trend", {})
    if qtc_trend.get("enabled"):
        trend_label = qtc_trend.get(
            "label",
            "For trend context only; clinical interpretation requires review.",
        )
        _add_key_value_table(doc, [("QTc Trend Note", trend_label)])

    # Section 4: Intra-Session Statistics
    _add_heading(doc, "Intra-Session Statistics")
    hr_vals = data.get("hr_values", [])
    rmssd_vals = data.get("rmssd_values", [])
    stats_rows = []
    if hr_vals:
        stats_rows.extend([
            ("HR Min", f"{min(hr_vals):.0f} bpm"),
            ("HR Max", f"{max(hr_vals):.0f} bpm"),
            ("HR Avg", f"{sum(hr_vals)/len(hr_vals):.0f} bpm"),
        ])
    else:
        stats_rows.append(("Heart Rate", "No data recorded"))
    if rmssd_vals:
        stats_rows.extend([
            ("RMSSD Min", f"{min(rmssd_vals):.1f} ms"),
            ("RMSSD Max", f"{max(rmssd_vals):.1f} ms"),
            ("RMSSD Avg", f"{sum(rmssd_vals)/len(rmssd_vals):.1f} ms"),
        ])
    else:
        stats_rows.append(("RMSSD", "No data recorded"))
    _add_key_value_table(doc, stats_rows, label_width_in=1.92, value_width_in=3.04)

    # Section 5: Session Visuals
    if visuals:
        visuals_heading = _add_heading(doc, "Session Visuals")
        # Avoid large whitespace gaps: keep only the minimum critical block together.
        visuals_heading.paragraph_format.keep_with_next = False
        ecg_path = visuals.get("ecg")
        if ecg_path and ecg_path.exists():
            _add_image_with_caption(
                doc,
                "Selected ECG strip (latest)",
                ecg_path,
                keep_block=False,
            )
        trend_path = visuals.get("trend")
        if trend_path and trend_path.exists():
            _add_image_with_caption(
                doc,
                "HR and RMSSD trend",
                trend_path,
                keep_block=True,
            )

    # Section 6: Annotations
    annotations = data.get("annotations", [])
    if annotations:
        _add_heading(doc, "Session Annotations")
        table = doc.add_table(rows=1 + len(annotations), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        ann_time_w = Inches(1.19)
        ann_text_w = Inches(2.94)
        table.rows[0].cells[0].text = "Time"
        table.rows[0].cells[1].text = "Annotation"
        table.rows[0].cells[0].width = ann_time_w
        table.rows[0].cells[1].width = ann_text_w
        for r_p in table.rows[0].cells:
            for p in r_p.paragraphs:
                for run in p.runs:
                    run.bold = True
        for r, (ts, txt) in enumerate(annotations, start=1):
            table.rows[r].cells[0].text = ts
            table.rows[r].cells[1].text = txt
            table.rows[r].cells[0].width = ann_time_w
            table.rows[r].cells[1].width = ann_text_w
        doc.add_paragraph("")

    # Section 7: Notes
    notes = data.get("notes", "").strip()
    _add_heading(doc, "Notes")
    doc.add_paragraph(notes if notes else "(No notes entered.)")
    if report_stage == "draft":
        draft_note = doc.add_paragraph(
            "This document is a draft exported before session finalization."
        )
        for run in draft_note.runs:
            run.bold = True

    # Section 8: Legal Disclaimer
    disclaimer = data.get("disclaimer", {}) or {}
    if disclaimer:
        _add_heading(doc, "Legal Disclaimer")
        warning = str(disclaimer.get("warning", "")).strip()
        if warning:
            warning_run = doc.add_paragraph().add_run(warning)
            warning_run.bold = True
        acknowledged_at = disclaimer.get("acknowledged_at") or "--"
        disclaimer_rows = [
            ("Acknowledgment Mode", str(disclaimer.get("acknowledgment_mode", "--"))),
            ("Acknowledged At", _report_datetime(acknowledged_at)),
            ("Source", str(disclaimer.get("source_path", "--"))),
        ]
        _add_key_value_table(doc, disclaimer_rows, label_width_in=1.81, value_width_in=2.32)
        disclaimer_text = str(disclaimer.get("text", "")).strip()
        if disclaimer_text:
            _add_heading(doc, "Saved Disclaimer Text", level=3)
            for line in disclaimer_text.splitlines():
                doc.add_paragraph(line if line.strip() else "")

    # Footer
    section = doc.sections[0]
    page_footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    page_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    page_footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    timestamp_run = page_footer.add_run(
        f"Generated by Hertz & Hearts on {_report_datetime(datetime.now())}"
    )
    timestamp_run.font.size = Pt(8)
    timestamp_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    page_footer.add_run("\t")
    _append_page_field(page_footer)

    doc.save(str(output))
